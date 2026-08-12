"""PyGeek 2026 심사용 DOCX 빌더.

제출 폼이 PDF와 **편집가능 원본**을 함께 요구하는데(.tex 는 허용 목록에 없음),
학회 배포 템플릿 `PyGeek_Template_Review.docx` 의 서식을 그대로 물려받아
`pygeek2026-en.tex` 의 본문을 한 글자도 바꾸지 않고 옮긴다.

템플릿은 이미 [1단 제목블록] + [2단 본문(continuous)] 구조라, 단 구성을 새로
만들지 않고 기존 구역 구분 문단만 남긴 채 내용만 교체한다.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[3]
TEMPLATE = ROOT / "docs/reference/pygeek-templates/PyGeek_Template_Review.docx"
FIGDIR = ROOT / "docs/figures"
OUT = Path(__file__).resolve().parent / "PyGeek2026_EN_심사용_익명.docx"

FONT = "Times New Roman"
BODY_PT, HEAD_PT, SMALL_PT, TABLE_PT = 10, 12, 8, 7
COL_W = 8.08  # (21.0 - 1.78*2 - 1.27) / 2 cm


# --------------------------------------------------------------------------
# LaTeX -> 서식 있는 run 변환
# --------------------------------------------------------------------------
TEXT_SUBS = [
    (r"\\url\{([^}]*)\}", r"\1"),
    (r"``", "\u201c"), (r"''", "\u201d"),
    (r"---", "\u2014"), (r"--", "\u2013"),
    (r"\\%", "%"), (r"\\&", "&"), (r"\\,", " "), (r"~", " "),
    (r"\\dots", "\u2026"),
]

# 수식 안에서 그대로 로마자로 두는 토큰
MATH_LITERAL = {
    r"\rightarrow": "\u2192",
    r"\pm": "\u00b1",
    r"\times": "\u00d7",
    r"\le": "\u2264",
    r"\ge": "\u2265",
}


def _plain(s: str) -> str:
    for pat, rep in TEXT_SUBS:
        s = re.sub(pat, rep, s)
    return s


def _emit_math(par, expr, size, bold):
    """수식 조각을 이탤릭 변수 / 로마자 숫자·기호 run 으로 나눠 붙인다."""
    i = 0
    while i < len(expr):
        ch = expr[i]

        if ch == "\\":
            m = re.match(r"\\text\{([^}]*)\}", expr[i:])
            if m:
                _run(par, m.group(1), size, bold, italic=False)
                i += m.end()
                continue
            m = re.match(r"\\[a-zA-Z]+", expr[i:])
            if m and m.group(0) in MATH_LITERAL:
                _run(par, MATH_LITERAL[m.group(0)], size, bold, italic=False)
                i += m.end()
                continue
            i += 1 if not m else m.end()
            continue

        if ch == "_":
            nxt = expr[i + 1:i + 2]
            if nxt == "{":
                m = re.match(r"_\{([^}]*)\}", expr[i:])
                sub, step = m.group(1), m.end()
            else:
                sub, step = nxt, 2
            _run(par, sub, size, bold, italic=True, subscript=True)
            i += step
            continue

        if ch.isalpha():
            m = re.match(r"[A-Za-z]+", expr[i:])
            _run(par, m.group(0), size, bold, italic=True)
            i += m.end()
            continue

        # 숫자 앞의 하이픈은 진짜 마이너스 기호로
        if ch == "-":
            _run(par, "\u2212", size, bold, italic=False)
            i += 1
            continue

        _run(par, ch, size, bold, italic=False)
        i += 1


def _run(par, text, size, bold, italic=False, subscript=False, superscript=False):
    if not text:
        return
    r = par.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if subscript:
        r.font.subscript = True
    if superscript:
        r.font.superscript = True
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return r


def _emit_text(par, chunk, size, bold):
    """일반 텍스트 조각 — 저자 소속 표시에 쓰이는 위첨자만 따로 떼어낸다."""
    for k, part in enumerate(re.split(r"\\textsuperscript\{([^}]*)\}", chunk)):
        _run(par, part if k % 2 else _plain(part), size, bold,
             superscript=bool(k % 2))


def rich(par, latex, size=BODY_PT, bold=False):
    """`$...$` 를 수식으로, 나머지는 일반 텍스트로 문단에 채운다."""
    for k, chunk in enumerate(re.split(r"\$([^$]*)\$", latex)):
        if k % 2 == 1:
            _emit_math(par, chunk, size, bold)
        else:
            _emit_text(par, chunk, size, bold)
    return par


# --------------------------------------------------------------------------
# 문단 헬퍼
# --------------------------------------------------------------------------
def style(par, *, align=None, before=0, after=0, line=1.15,
          indent=None, hanging=None, keep=False):
    pf = par.paragraph_format
    if align is not None:
        par.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent is not None:
        pf.left_indent = Cm(indent)
    if hanging is not None:
        pf.first_line_indent = Cm(-hanging)
    pf.keep_with_next = keep
    return par


def hrule(par):
    """제목블록 위아래 가로줄 (LaTeX 판의 \\rule 대응)."""
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    par._p.get_or_add_pPr().append(pbdr)
    return par


TWIPS_PER_CM = 566.93


def fix_widths(tbl, widths_cm):
    """Word 의 자동 너비는 2단 폭을 넘겨 표를 잘라먹는다.

    고정 레이아웃에서 실제로 참조되는 것은 tblGrid 이므로, 셀 너비만이 아니라
    그리드 자체를 못박아야 한다.
    """
    tbl.autofit = False                       # w:tblLayout = fixed
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for col, cm in zip(grid.findall(qn("w:gridCol")), widths_cm):
        col.set(qn("w:w"), str(int(cm * TWIPS_PER_CM)))
    for ci, cm in enumerate(widths_cm):
        for row in tbl.rows:
            row.cells[ci].width = Cm(cm)


def tight_margins(tbl, dxa=45):
    """기본 셀 여백(양쪽 0.19cm)은 좁은 2단 표에서 글자 자리를 크게 잡아먹는다."""
    for row in tbl.rows:
        for cell in row.cells:
            mar = OxmlElement("w:tcMar")
            for side in ("left", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:w"), str(dxa))
                el.set(qn("w:type"), "dxa")
                mar.append(el)
            cell._tc.get_or_add_tcPr().append(mar)


def set_borders(tbl, rows_top, rows_bottom):
    """booktabs 풍 — 지정한 행에만 가로선."""
    for idx, row in enumerate(tbl.rows):
        for cell in row.cells:
            tcpr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge, want in (("top", idx in rows_top), ("bottom", idx in rows_bottom)):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single" if want else "nil")
                el.set(qn("w:sz"), "8" if want else "0")
                el.set(qn("w:color"), "000000")
                borders.append(el)
            tcpr.append(borders)


# --------------------------------------------------------------------------
# 본문 내용 (pygeek2026-en.tex 원문 그대로)
# --------------------------------------------------------------------------
TITLE = ("A Factor Decomposition Method for Separating Store Substitution "
         "and Embedding Placement Effects in Vector Database Migration")

# 심사본은 double-blind라 저자 정보가 없다. 카메라레디 빌더가 이 넷을 채우면
# 제목 블록과 본문 끝에 해당 절이 생긴다 (build_docx_camera.py 참고).
AUTHORS = None          # 저자명 한 줄
AFFILIATIONS = []       # 소속 줄들
CORRESPONDING = None    # 교신저자 줄
COI = None              # Conflict of Interest 본문

# 템플릿 실측: 저자명 8pt bold(소속 10pt 보다 작다), `Keywords:` 라벨 12pt bold.
AUTHOR_PT = SMALL_PT
KEYWORD_LABEL_PT = 9

ABSTRACT = r"""When the vector layer of a retrieval-augmented generation (RAG) system is migrated to a different database, the effect is usually assessed by an end-to-end comparison of the system before and after the migration. That comparison, however, blends into a single measurement two changes that are independent yet occur together in practice, namely substituting the store engine and relocating where the embedding is computed, which makes it impossible to tell which factor caused the change. This paper proposes a method that introduces a bridge condition, in which the target store is loaded with the embeddings the source system has already produced, so that one comparison in which the two factors are entangled is split into two comparisons that each differ in one factor only. The bridge condition requires no re-embedding, so it is inexpensive and leaves the running service untouched. Applied to a Korean product-review RAG service in production (1,347 chunks, 43 queries), no change in retrieval quality was detected for the store substitution (top-10 Jaccard 0.971, relevance difference $-$0.042, $p = 0.52$), whereas relocating the embedding into the database produced a significant, medium-sized effect ($p = 0.003$, Cohen's $d_z = 0.47$). Virtually all of the observed change therefore originates from the embedding placement factor. The method does not judge the superiority of any particular product and applies to migrations to converged databases in general."""

KEYWORDS = ("Retrieval-Augmented Generation, Vector Database, "
            "In-Database Embedding, Migration, Factor Decomposition")

INTRO = [
    r"""A retrieval-augmented generation (RAG) system consists of an embedding model that turns queries and documents into vectors, a store that indexes and searches those vectors, and a language model that generates an answer grounded in the retrieved results [1], [2]. Among these, the component changed most often after deployment is the vector layer. A representative path leads from a general-purpose relational database with a vector extension to a converged database that has a vector type and an inference runtime inside it, that is, one in which the database itself performs the embedding computation.""",
    r"""The effect of such a migration is usually assessed by an end-to-end comparison that weighs the whole system before and after the change at once. In this comparison, however, two factors change at the same time. The first is the store, that is, the engine that holds the vectors and the indexing method it uses. The second is the placement, that is, whether the embedding is computed by an external service or by an inference runtime inside the database. When a migration is motivated by governance, moving the embedding inside as well is the preferred configuration, and a database that hosts a model limits the size of the model file. This is why the two factors change together in practice.""",
    r"""As a result, the end-to-end difference alone does not reveal which factor caused the change. When quality drops, it cannot be told apart whether the store or the model is at fault, and when nothing changes, it cannot be known whether the two factors cancelled each other out. In either case the conclusion cannot be carried over as is to an environment with a different instance class or model-size constraint.""",
    r"""This paper inserts an intermediate condition, the bridge condition, that loads the target store with the embeddings the source system has already produced, and thereby splits one comparison in which the two factors are entangled into two comparisons that each differ in one factor at a time. We apply the proposed method to a Korean review-analysis service migrated from a PostgreSQL vector extension to a converged database that supports in-database embedding, and measure it on 1,347 review chunks from the production corpus and 43 fixed queries. This paper does not judge which product is superior, and the latency figures reflect the measurement environment rather than the products.""",
]

SEC21_INTRO = r"""The target system retrieves Korean product reviews and generates answers that cite the retrieved evidence. Both architectures operate in four stages: query input, query embedding, vector search, and answer generation. Retrieval is a hybrid scheme that combines vector similarity with metadata filters and source-based prioritization, and this rule and its parameters were kept identical across the three conditions. Because both are expressed as a single SQL statement, the store and the embedding placement can be swapped while the application logic is left as it is."""

FIG1_TEXT = r"""Fig. 1 shows the Oracle-based architecture. The application sends a single SQL statement carrying the query string (step 1), and the database converts the query into a 384-dimensional vector with its built-in ONNX inference runtime (step 2) and then, within the same SQL statement, performs an exact search over the VECTOR column together with metadata filtering (step 3). The language model cites the returned top-$k$ chunks to generate an answer (step 4)."""

FIG2_TEXT = r"""Fig. 2 shows the PostgreSQL-based architecture. The application calls an external embedding service to convert the query into a 1024-dimensional vector (step 2) and then requests the search with that vector as an argument (step 3). The store explores candidates with an HNSW approximate index and applies the same filters, and the remaining steps are as in Fig. 1."""

SEC21_END = r"""The two architectures differ in two things at once, the store and the embedding placement. One is which engine holds the vectors and with what indexing method, and the other is whether the embedding is computed outside or inside the database. An end-to-end comparison alone cannot tell which of the two produced the difference. The next two subsections summarize the characteristics of each architecture from these two viewpoints, and Section 3 presents a way to separate them."""

SEC22 = r"""First, it provides a native VECTOR type and a built-in ONNX inference runtime, so embedding and search are combined into a single SQL statement [3]. Second, it performs an exact search that compares every vector, without an approximate index that narrows candidates in advance, so no correct result is missed because of the index. Third, the network boundary is reduced to one, and because the query and the retrieved results stay inside the database, controls such as encryption and auditing apply to the entire retrieval path. Fourth, a built-in model follows the model-file size limit the database allows, and in this environment the largest multilingual model within that limit was multilingual-e5-small (384 dimensions) [4]."""

SEC23 = r"""First, it provides vector search through an extension module (pgvector) and contains no inference runtime [5]. The embedding model can be chosen separately from the store, so there is no size constraint, and this architecture uses the 1024-dimensional BGE-M3 [6]. Second, it uses an HNSW approximate index that narrows candidates as it searches, so search time grows gently as the data grows, at the cost of missing some correct results [7]. Third, because the embedding is computed externally, the query text leaves the store boundary and one more network round trip is required."""

SEC31 = [
    r"""In this experiment we denote a retrieval configuration by the triple $(S, E, P)$, where $S$ is the store, $E$ is the embedding model, and $P$ is where the embedding is executed. Comparing the two architectures of Section 2 as they are amounts to comparing $(S_0, E_0, \text{external})$ with $(S_1, E_1, \text{in-database})$; all three coordinates differ, so the contribution of each factor cannot be separated. We therefore add $(S_1, E_0, \text{external})$, a bridge condition that uses the target store but is loaded with the vectors the source system produced, and the three conditions are summarized in Table 1. The bridge condition needs no recomputation of the vectors. Because the columns in production are left as they are and only one more vector column is created and filled, its preparation cost is a single bulk load and the running service is not touched.""",
    r"""In condition C, however, moving the placement also changes the model, because the file-size limit of Section 2.2 makes it impossible to host the source model inside the database. Choosing in-database placement means accepting that limit as well, so this coupling was left in place on purpose. Separating the two would require one more condition, $(S_1, E_1, \text{external})$, which runs the smaller model outside the database, and we leave this to future work.""",
    r"""The corpus consists of 1,347 review chunks from a single product domain of the production service, loaded identically into both stores after verifying in advance that the chunk identifiers correspond one to one. The query set was fixed at 43 Korean queries, 13 drawn from production logs and 30 synthetic queries written so as to cover evenly the evaluation aspects and purchase-journey stages the service handles.""",
    r"""Three metrics are measured. The first is the agreement of the retrieved results: how much the top-10 results the two conditions return for the same query overlap (Jaccard), how similar their rankings are (Spearman rank correlation), and, taking the exact search of the bridge condition as ground truth, how many of those the baseline recovers (Recall@10). The second is the relevance of the retrieved evidence: a large language model scores the top-5 chunks on a 0--2 scale [8], and 215 judgments per condition are averaged per query and then compared in pairs across the same queries. The third is the actual time taken to process a single query.""",
    r"""The target database ran on the lowest-tier free instance (Always Free, 2 GB RAM) and the external embedding service on a local laptop (Apple M5 Pro). Confidence intervals were obtained by bootstrap resampling of the measurements 10,000 times.""",
]

SEC32 = [
    r"""The measurements by factor are summarized in Table 2.""",
    r"""Looking first at ranking agreement, the two stores loaded with identical vectors returned exactly the same top-10 set for 37 of the 43 queries, and even the minimum was 0.667. The results the baseline missed because of its approximate index amount to about 1.6\%.""",
    r"""For the store factor, the paired difference in relevance scores across the same queries was not statistically significant and the effect size was negligible ($t(42) = -0.643$). A Wilcoxon test on the 28 queries that remain after excluding those with identical scores leads to the same conclusion ($p = 0.908$).""",
    r"""Since ``no difference'' is itself our claim, we did not rely only on a non-significant $p$ value but tested equivalence separately [9]. Two one-sided tests examining whether the difference falls within $\pm$0.15 points gave $p = 0.052$: the lower bound of the confidence interval, $-$0.172, falls outside this range, so the significance level was not met. We therefore report not that equivalence has been established but that no effect was detected and that its magnitude is bounded to a small range. Since the direction in which the interval extends is the one where the target store scores higher, there is no basis for concluding that substituting the store degraded quality.""",
    r"""By contrast, between B and C, which share the same store and the same search method, the results differ greatly. The top-10 Jaccard falls to 0.175, so mostly different evidence is retrieved for the same query, and the relevance score decreases from 1.270 to 1.014. This difference is significant with a medium effect size ($t(42) = 3.104$), and a Wilcoxon test on the 38 queries that remain after excluding those with identical scores leads to the same conclusion ($p = 0.006$). Because no store effect was detected in the change from A to B, virtually all of the end-to-end quality change is due to the placement factor and the model-capacity limit that comes with it. Had only an end-to-end comparison been performed, the roughly 20\% drop in relevance would have been blamed on the store, and reverting the store, the remedy that follows from that diagnosis, would not bring the quality back.""",
    r"""For search latency, the ratio of C to B was 8.27 times at the median (Wilcoxon $p < 0.001$). The increase in C is not a property of in-database embedding itself but a consequence of running inference on a free-tier instance, so it cannot be generalized to higher-tier hardware. Since A and B produce the same vectors with the same external service, the query embedding time was measured once and applied to both conditions, and the latency difference between them therefore reflects only the store search time.""",
    r"""Finally, because 30 of the 43 queries are synthetic, we checked whether they behave differently from the 13 real queries. Relevance scores were systematically lower on the synthetic subset (A: 1.147 vs. 1.415, $p = 0.045$; C: 0.880 vs. 1.323, $p = 0.007$; Mann--Whitney), so the absolute level of relevance must not be read as service quality. In contrast, ranking agreement, the basis on which the cause is attributed, showed no significant difference between the two subsets ($p = 0.441$, $p = 0.570$). Absolute scores depend on where the queries come from, but the conclusion about the cause is the same on both.""",
]

CONCLUSION = [
    r"""Adding a single bridge condition makes it possible to split an end-to-end comparison into two single-factor comparisons. Applied to a Korean review RAG service in production, no effect of the store factor was detected, and the entire observed change originated from the placement factor. Changing the placement is a trade in governance rather than in performance. In exchange for keeping the query and the review text inside the database boundary, one has to accept the capacity limit of the model and the inference performance of the database tier. Therefore, if governance of the stored data is the only goal, replacing the store alone is enough, and if the query path must also stay inside the boundary, one should first confirm that the largest model that can be hosted is sufficient for that domain. The bridge condition makes that confirmation possible before the migration is carried out.""",
    r"""The limitations are as follows. The study targets a single Korean corpus of 1,347 chunks, human assessors were not involved in the relevance scoring, and the measurements were taken on an entry-level instance. In addition, as stated in Section 3.1, placement and model capacity are coupled by construction, and the result for the store factor does not establish equivalence.""",
]

TABLE1 = [
    ["", "A (baseline)", "B (bridge)", "C (migrated)"],
    ["Store", "PostgreSQL + pgvector", "Oracle 26ai", "Oracle 26ai"],
    ["Embedding model", "BGE-M3 (1024)", "BGE-M3 (1024)", "e5-small (384)"],
    ["Placement", "external", "external", "in-database"],
    ["Vector search", "HNSW approx.", "exact", "exact"],
]
TABLE1_NOTE = r"""A$\rightarrow$B changes only the store factor, and B$\rightarrow$C only the placement factor. e5-small denotes multilingual-e5-small."""

TABLE2 = [
    ["Metric", r"A $\rightarrow$ B (store)", r"B $\rightarrow$ C (placement)"],
    [r"Top-10 Jaccard [95\% CI]", "0.971 [0.946, 0.992]", "0.175 [0.133, 0.221]"],
    ["Spearman rank correlation", "0.958", "---"],
    ["Recall@10 (vs. exact search)", "0.984", "---"],
    ["Relevance score (0--2)", r"1.228 $\rightarrow$ 1.270", r"1.270 $\rightarrow$ 1.014"],
    ["Paired difference [95\\% CI]", r"$-$0.042 [$-$0.172, 0.084]", "0.256 [0.098, 0.423]"],
    [r"Significance / effect size $d_z$", r"$p = 0.523$ / $-$0.098", "$p = 0.003$ / 0.473"],
    ["Mean search latency (ms)", r"150.2 $\rightarrow$ 110.0", r"110.0 $\rightarrow$ 947.0"],
    ["External calls / boundaries", r"1$\rightarrow$1 / 2$\rightarrow$2", r"1$\rightarrow$0 / 2$\rightarrow$1"],
]
TABLE2_NOTE = r"""Relevance is the per-query mean over the 43 queries (215 judgments per condition), and $p$ values are from a paired $t$-test."""

REFERENCES = [
    r"""[1] P. Lewis et al., ``Retrieval-augmented generation for knowledge-intensive NLP tasks,'' Advances in Neural Information Processing Systems, vol. 33, pp. 9459-9474, 2020.""",
    r"""[2] Y. Gao et al., ``Retrieval-augmented generation for large language models: A survey,'' arXiv preprint arXiv:2312.10997, 2023.""",
    r"""[3] Oracle, ``Oracle AI Vector Search User's Guide,'' Oracle AI Database 26ai documentation, Available: \url{https://docs.oracle.com/en/database/oracle/oracle-database/26/vecse/}, 2026, [Accessed: Aug. 8, 2026].""",
    r"""[4] L. Wang, N. Yang, X. Huang, L. Yang, R. Majumder, and F. Wei, ``Multilingual E5 text embeddings: A technical report,'' arXiv preprint arXiv:2402.05672, 2024.""",
    r"""[5] pgvector, ``pgvector: Open-source vector similarity search for Postgres,'' GitHub repository, Available: \url{https://github.com/pgvector/pgvector}, 2026, [Accessed: Aug. 8, 2026].""",
    r"""[6] J. Chen, S. Xiao, P. Zhang, K. Luo, D. Lian, and Z. Liu, ``M3-Embedding: Multi-linguality, multi-functionality, multi-granularity text embeddings through self-knowledge distillation,'' Findings of the Association for Computational Linguistics: ACL 2024, Bangkok, Thailand, pp. 2318-2335, \url{https://doi.org/10.18653/v1/2024.findings-acl.137}, Aug., 2024.""",
    r"""[7] Y. A. Malkov and D. A. Yashunin, ``Efficient and robust approximate nearest neighbor search using hierarchical navigable small world graphs,'' IEEE Transactions on Pattern Analysis and Machine Intelligence, vol. 42, no. 4, pp. 824-836, \url{https://doi.org/10.1109/TPAMI.2018.2889473}, 2020.""",
    r"""[8] L. Zheng et al., ``Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena,'' Advances in Neural Information Processing Systems 36 (NeurIPS 2023), Datasets and Benchmarks Track, 2023.""",
    r"""[9] D. Lakens, ``Equivalence tests: A practical primer for t tests, correlations, and meta-analyses,'' Social Psychological and Personality Science, vol. 8, no. 4, pp. 355-362, \url{https://doi.org/10.1177/1948550617697177}, 2017.""",
]


# --------------------------------------------------------------------------
# 빌드
# --------------------------------------------------------------------------
def build():
    doc = Document(str(TEMPLATE))
    body = doc.element.body

    # 템플릿의 [1단 제목블록 | 2단 본문] 구역 구분 문단만 남기고 내용 비우기
    break_p = None
    for el in body.iterchildren():
        if el.tag == qn("w:p") and el.find(f".//{qn('w:sectPr')}") is not None:
            break_p = el
            break
    assert break_p is not None, "구역 구분 문단을 찾지 못했다"
    for el in list(body.iterchildren()):
        if el is not break_p and el.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(el)

    # 머리글의 연도 자리표시자를 실제 연도로
    for section in doc.sections:
        for hdr in (section.header, section.even_page_header):
            for par in hdr.paragraphs:
                if "20xx" not in par.text:
                    continue
                merged = par.text.replace("20xx", "2026")
                for extra in par.runs[1:]:
                    extra._element.getparent().remove(extra._element)
                par.runs[0].text = merged

    def before_break():
        p = OxmlElement("w:p")
        break_p.addprevious(p)
        return Paragraph(p, doc._body)

    # ---------- 1단 제목 블록 ----------
    style(rich(before_break(), TITLE, size=16, bold=True),
          align=WD_ALIGN_PARAGRAPH.CENTER, after=5, line=1.0)
    if AUTHORS:
        style(rich(before_break(), AUTHORS, size=AUTHOR_PT, bold=True),
              align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=4, line=1.0)
        for aff in AFFILIATIONS:
            style(rich(before_break(), aff, size=BODY_PT), after=2, line=1.0)
        style(rich(before_break(), CORRESPONDING, size=BODY_PT),
              before=3, after=4, line=1.0)
    hrule(style(before_break(), after=3, line=1.0))
    style(rich(before_break(), "Abstract", size=HEAD_PT, bold=True), after=3, keep=True)
    style(rich(before_break(), ABSTRACT, size=SMALL_PT),
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=4)
    kw = before_break()
    rich(kw, "Keywords: ", size=KEYWORD_LABEL_PT, bold=True)
    rich(kw, KEYWORDS, size=SMALL_PT)
    style(kw, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=4)
    hrule(style(before_break(), after=6, line=1.0))

    # ---------- 2단 본문 ----------
    def head(text, level=1):
        p = style(rich(doc.add_paragraph(), text, size=HEAD_PT, bold=True),
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  before=5 if level == 1 else 4, after=2, keep=True)
        return p

    def para(text):
        return style(rich(doc.add_paragraph(), text),
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3)

    def figure(png, caption):
        p = doc.add_paragraph()
        p.add_run().add_picture(str(png), width=Cm(COL_W - 0.2))
        style(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2, keep=True)
        style(rich(doc.add_paragraph(), caption, size=SMALL_PT),
              align=WD_ALIGN_PARAGRAPH.CENTER, after=5)

    def table(rows, title, note, widths):
        style(rich(doc.add_paragraph(), title, size=SMALL_PT),
              align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2, keep=True)
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        fix_widths(tbl, widths)
        for ri, row in enumerate(rows):
            for ci, cell_text in enumerate(row):
                cell = tbl.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                rich(p, cell_text, size=TABLE_PT, bold=(ri == 0))
                style(p, after=0, line=1.0,
                      align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                      else WD_ALIGN_PARAGRAPH.CENTER)
        set_borders(tbl, rows_top={0}, rows_bottom={0, len(rows) - 1})
        tight_margins(tbl)
        style(rich(doc.add_paragraph(), note, size=TABLE_PT),
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=3, after=5)

    head("1. Introduction")
    for t in INTRO:
        para(t)

    head("2. Proposed Method")
    head("2.1 System Architecture", 2)
    para(SEC21_INTRO)
    para(FIG1_TEXT)
    figure(FIGDIR / "paper-fig1-arch-oracle.png", "Fig. 1. Oracle-based system architecture")
    para(FIG2_TEXT)
    figure(FIGDIR / "paper-fig2-arch-pgvector.png", "Fig. 2. PostgreSQL-based system architecture")
    para(SEC21_END)

    head("2.2 Characteristics of the Oracle-Based Architecture", 2)
    para(SEC22)
    head("2.3 Characteristics of the PostgreSQL-Based Architecture", 2)
    para(SEC23)

    head("3. Performance Evaluation")
    head("3.1 Experimental Setup", 2)
    para(SEC31[0])
    table(TABLE1, "Table 1. Experimental conditions", TABLE1_NOTE,
          [1.90, 2.06, 2.06, 2.06])
    for t in SEC31[1:]:
        para(t)

    head("3.2 Experimental Results", 2)
    para(SEC32[0])
    table(TABLE2, "Table 2. Results by factor", TABLE2_NOTE,
          [3.40, 2.34, 2.34])
    for t in SEC32[1:]:
        para(t)

    head("4. Conclusion")
    for t in CONCLUSION:
        para(t)

    if COI:
        style(rich(doc.add_paragraph(), "Conflict of Interest",
                   size=HEAD_PT, bold=True),
              align=WD_ALIGN_PARAGRAPH.CENTER, before=5, after=2, keep=True)
        style(rich(doc.add_paragraph(), COI),
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3)

    style(rich(doc.add_paragraph(), "References", size=HEAD_PT, bold=True),
          align=WD_ALIGN_PARAGRAPH.CENTER, before=5, after=2, keep=True)
    for ref in REFERENCES:
        # 양쪽 정렬은 긴 URL 때문에 단어 간격이 크게 벌어진다 (LaTeX 판도 \raggedright)
        style(rich(doc.add_paragraph(), ref, size=SMALL_PT),
              align=WD_ALIGN_PARAGRAPH.LEFT, after=0, line=1.15,
              indent=0.5, hanging=0.5)

    doc.save(str(OUT))
    print(f"저장: {OUT}")


if __name__ == "__main__":
    build()
