"""DOCX 가 제출용 PDF 와 한 글자도 다르지 않은지 검증한다.

빌드 스크립트에 본문을 하드코딩했으므로 옮겨 적는 과정의 오타·누락이
유일한 실질 위험이다. PDF 는 .tex 를 컴파일한 결과이므로 원문의 렌더링
그 자체다. 따라서 두 파일의 텍스트 흐름을 문자 단위로 대조한다.

PDF 는 양쪽 정렬 과정에서 단어를 하이픈으로 쪼개므로, 비교 전에 공백과
하이픈을 모두 제거해 그 차이를 무력화한다(하이픈은 양쪽에서 똑같이 지운다).
"""

import re
import subprocess
import sys
import unicodedata
from collections import Counter
from pathlib import Path

from docx import Document

HERE = Path(__file__).resolve().parent

# `--camera` 를 주면 카메라레디 판을 검증한다. 두 판은 익명성 요구가 정반대라
# (심사본=저자 없어야 함, 카메라레디=저자 있어야 함) 검사도 갈라진다.
CAMERA = "--camera" in sys.argv
STEM = "PyGeek2026_EN_카메라레디_최종" if CAMERA else "PyGeek2026_EN_심사용_익명"
DOCX = HERE / f"{STEM}.docx"
PDF = HERE / f"{STEM}.pdf"

# 머리글·바닥글 — PDF 에는 텍스트로 박혀 나오지만 DOCX 본문에는 없다
CHROME = [
    "2026 International Conference on PyGeek (PyGeek 2026)",
    "Vol. xx, 2026",
    "ISSN (Online): 0000-0000",
]


def canon(s: str) -> str:
    """비교용 정규화 — 유니코드 변형·공백·하이픈·따옴표 차이를 제거."""
    s = unicodedata.normalize("NFKC", s)
    for a, b in [("−", "-"), ("–", "-"), ("—", "-"), ("‐", "-"),
                 ("“", '"'), ("”", '"'), ("‘", "'"), ("’", "'"),
                 ("→", ">"), ("±", "+-"), ("－", "-")]:
        s = s.replace(a, b)
    s = s.lower()
    return re.sub(r"[\s\-]+", "", s)


def docx_text() -> str:
    d = Document(str(DOCX))
    parts = []

    def walk(container):
        for block in container.element.body.iterchildren() if hasattr(container, "element") else []:
            pass

    for p in d.paragraphs:
        parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def docx_flow() -> str:
    """문서 순서대로 문단·표를 훑는다 (표가 본문 사이에 있으므로)."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    d = Document(str(DOCX))
    out = []
    for el in d.element.body.iterchildren():
        if el.tag == qn("w:p"):
            out.append(Paragraph(el, d).text)
        elif el.tag == qn("w:tbl"):
            for row in Table(el, d).rows:
                out.extend(c.text for c in row.cells)
    return "\n".join(out)


def pdf_text() -> str:
    """PDF 본문 텍스트 — 머리글·바닥글은 DOCX 본문에 없으므로 걷어낸다.

    쪽번호를 '홀로 남은 숫자 줄' 로 지우면 안 된다. 카메라레디의 소속 위첨자
    번호도 똑같이 홀로 있는 줄로 떨어져 나와 같이 지워지기 때문이다. 대신
    페이지를 나눈 뒤 그 페이지의 번호만 끝에서 하나 떼어낸다.
    """
    raw = subprocess.run(["pdftotext", str(PDF), "-"],
                         capture_output=True, text=True, check=True).stdout
    for c in CHROME:
        raw = raw.replace(c, " ")
    out = []
    for i, page in enumerate(raw.split("\f"), 1):
        lines = page.rstrip().splitlines()
        # 쪽번호는 페이지 맨 아래에 있다. 위에서부터 지우면 소속 위첨자를 먼저 만난다.
        # 2단 조판이라 쪽번호가 본문 단어에 붙어 나오기도 한다("...generation. Re1"),
        # 그래서 단어 경계(\b)를 쓰지 않고 줄 끝의 해당 숫자만 떼어낸다.
        for j in range(len(lines) - 1, -1, -1):
            if lines[j].strip():
                lines[j] = re.sub(rf"{i}\s*$", "", lines[j])
                break
        out.append("\n".join(lines))
    return "\n".join(out)


def docx_blocks():
    """DOCX 를 비교 단위(문단 / 표의 한 행)로 쪼갠다."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    d = Document(str(DOCX))
    for el in d.element.body.iterchildren():
        if el.tag == qn("w:p"):
            yield Paragraph(el, d).text
        elif el.tag == qn("w:tbl"):
            for row in Table(el, d).rows:
                yield "".join(c.text for c in row.cells)


def split_head(s: str):
    """제목 블록(제목·저자·소속)과 그 이후 본문을 나눈다.

    카메라레디의 소속 위첨자 번호는 PDF 에서 '1' 만 있는 줄로 떨어져 나오는데,
    쪽번호를 걸러내는 규칙이 그것까지 지워 버린다. 본문 대조에서 제목 블록을
    떼어내고 저자 정보는 따로 확인하는 편이 정확하다.
    """
    i = s.find("Abstract")
    return (s[:i], s[i:]) if i != -1 else ("", s)


def main() -> int:
    dx_full, pdf_full = docx_flow(), pdf_text()

    if CAMERA:
        dx_head, dx_raw = split_head(dx_full)
        pdf_head, pdf_raw = split_head(pdf_full)
        # 제목 블록은 저자·소속 문자열 존재 여부로 따로 검사한다
        ch, cp = canon(dx_head), canon(pdf_head)
        same = ch == cp
        print(f"{'✅' if same else '❌'} 제목 블록(제목·저자·소속) "
              f"{'DOCX=PDF 일치' if same else '불일치'}")
        if not same:
            print(f"   DOCX만: {(Counter(ch)-Counter(cp)).most_common(8)}")
            print(f"   PDF만 : {(Counter(cp)-Counter(ch)).most_common(8)}")
    else:
        dx_raw, pdf_raw = dx_full, pdf_full

    a, b = canon(dx_raw), canon(pdf_raw)
    ok = True

    # 1) 문자 다중집합 — 글자 하나라도 늘거나 줄면 걸린다
    print(f"DOCX 문자수 {len(a):,} / PDF 문자수 {len(b):,}")
    if len(a) == len(b) and Counter(a) == Counter(b):
        print("✅ 문자 다중집합 완전 일치 (하이픈·공백 무시)")
    else:
        ok = False
        print(f"❌ 문자 구성 불일치 — DOCX에만 {(Counter(a)-Counter(b)).most_common(10)} / "
              f"PDF에만 {(Counter(b)-Counter(a)).most_common(10)}")

    # 2) 어긋나는 구간이 '순수한 위치 이동'인지 확인한다.
    #    그림·표는 PDF 에서 페이지 상단으로 떠오르고 표는 열 방향으로 읽히므로 순서가 다르다.
    #    따라서 순서가 아니라, 어긋난 구간들이 양쪽에서 같은 문자로 이뤄졌는지를 본다.
    import difflib

    wa = re.findall(r"\S+", dx_raw.lower())
    wb = re.findall(r"\S+", pdf_raw.lower())
    sm = difflib.SequenceMatcher(None, wa, wb, autojunk=False)
    left, right, regions = [], [], 0
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        regions += 1
        left.append(" ".join(wa[i1:i2]))
        right.append(" ".join(wb[j1:j2]))
    cl, cr = Counter(canon(" ".join(left))), Counter(canon(" ".join(right)))
    if cl == cr:
        print(f"✅ 어긋난 {regions}개 구간이 모두 동일 문자로 구성 — 내용 변화 없이 배치만 다름")
    else:
        ok = False
        print(f"❌ 이동으로 설명되지 않는 차이 — DOCX쪽 {(cl-cr).most_common(12)} / "
              f"PDF쪽 {(cr-cl).most_common(12)}")

    # 수치 대조 — 논문의 생명
    pat = r"\d+(?:,\d{3})*(?:\.\d+)?"
    dn, pn = Counter(re.findall(pat, dx_raw)), Counter(re.findall(pat, pdf_raw))
    if dn == pn:
        print(f"✅ 수치 토큰 {sum(dn.values())}개 전부 일치")
    else:
        ok = False
        print(f"❌ 수치 불일치 — DOCX에만 {(dn - pn).most_common(15)} / "
              f"PDF에만 {(pn - dn).most_common(15)}")

    low = dx_full.lower()
    if CAMERA:
        # 카메라레디는 반대로 저자 정보가 **있어야** 한다
        need = ["junha yoon", "ohhyeon gwon", "junhaeng lee", "hyeji roh",
                "soowan cho", "suhee kim", "sangsoon lim",
                "sungkyul university", "chung-ang university", "deep insight lab",
                "corresponding author", "slim@cau.ac.kr", "conflict of interest"]
        missing = [w for w in need if w not in low]
        print(f"{'✅' if not missing else '❌'} 저자 정보·COI "
              f"{'전부 포함' if not missing else '누락: ' + str(missing)}")
        ok = ok and not missing

        # 지도교수 판단으로 빼기로 한 절이 실수로 들어가지 않았는지
        banned = ["acknowledg", "author information", "research interest"]
        stray = [w for w in banned if w in low]
        print(f"{'✅' if not stray else '❌'} 생략 대상 절 "
              f"{'없음' if not stray else '남아있음: ' + str(stray)}")
        ok = ok and not stray

        # 리뷰어 지적(용어 일관성) 반영 여부 — 옛 표현이 남으면 실패
        canon_txt = re.sub(r"\s+", " ", dx_full)
        # 첫 등장에서 정의한 뒤 재사용하는 약어 자체는 정상이므로,
        # 정의 없이 쓰던 '첫 등장' 문구만 골라 잔존 여부를 본다.
        old_terms = ["model-size constraint", "model-file size limit",
                     "file-size limit", "model-capacity limit",
                     "capacity limit of the model",
                     "where the embedding is executed", "$p = 0.52$",
                     "with its built-in ONNX inference runtime",
                     "candidates with an HNSW approximate index"]
        left = [t for t in old_terms if t in canon_txt]
        print(f"{'✅' if not left else '❌'} 용어 일관성 수정 "
              f"{'반영됨' if not left else '옛 표현 잔존: ' + str(left)}")
        ok = ok and not left
        n_unified = len(re.findall("model size limit", canon_txt))
        print(f"{'✅' if n_unified == 6 else '❌'} 'model size limit' {n_unified}회 (기대 6회)")
        ok = ok and n_unified == 6
    else:
        # 심사본은 double-blind — 저자를 식별할 무엇도 없어야 한다
        ident = ["yoon", "gwon", "roh", "suhee", "soowan", "junhaeng", "junha",
                 "sangsoon", "sungkyul", "chung-ang", "deep insight", "cau.ac.kr",
                 "dilab", "conflict of interest", "acknowledg"]
        hits = [w for w in ident if w in low]
        print(f"{'✅' if not hits else '❌'} 익명성·camera-ready 전용 절 "
              f"{'통과' if not hits else '위반: ' + str(hits)}")
        ok = ok and not hits

    # 서식 — 학회 템플릿에서 물려받았는지
    d = Document(str(DOCX))
    s0, s1 = d.sections[0], d.sections[1]
    checks = [
        ("A4 (21.0 x 29.7cm)", round(s0.page_width.cm, 1) == 21.0 and round(s0.page_height.cm, 1) == 29.7),
        ("여백 1.78cm", all(round(getattr(s0, f"{e}_margin").cm, 2) == 1.78
                            for e in ("left", "right", "top", "bottom"))),
        ("제목블록 1단", 'w:num' not in s0._sectPr.xml.split('<w:cols')[1].split('>')[0]),
        ("본문 2단 (연속)", 'w:num="2"' in s1._sectPr.xml and 'continuous' in s1._sectPr.xml),
        ("그림 2개 삽입", len(d.inline_shapes) == 2),
        ("표 2개", len(d.tables) == 2),
    ]
    for name, good in checks:
        print(f"{'✅' if good else '❌'} {name}")
        ok = ok and good

    # 글꼴·크기 — 텍스트 대조만으로는 서식 오류를 못 잡는다.
    # 기대값은 학회 템플릿 DOCX 에서 실측한 것이다.
    def sizes(par):
        return {r.font.size.pt for r in par.runs if r.font.size}

    def fonts(par):
        return {r.font.name for r in par.runs if r.font.name}

    spec = [("제목", 0, {16.0}, True)]
    if CAMERA:
        spec += [("저자명", 1, {8.0}, True),
                 ("소속 1행", 2, {10.0}, False),
                 ("교신저자", 7, {10.0}, False)]
    ps = [p for p in d.paragraphs if p.text.strip()]
    for label, idx, want, want_bold in spec:
        got = sizes(ps[idx])
        good = got == want and (not want_bold or all(r.bold for r in ps[idx].runs))
        print(f"{'✅' if good else '❌'} {label} {sorted(got)}pt "
              f"(기대 {sorted(want)}pt{', bold' if want_bold else ''})")
        ok = ok and good

    kw = next((p for p in ps if p.text.startswith("Keywords:")), None)
    if kw:
        lab = kw.runs[0]
        good = lab.font.size.pt == 12.0 and lab.bold
        print(f"{'✅' if good else '❌'} 'Keywords:' 라벨 {lab.font.size.pt}pt "
              f"bold={lab.bold} (기대 12.0pt bold)")
        ok = ok and good

    # 초록 분량 — 템플릿 규정 150~200 단어
    m = re.search(r"Abstract\s*\n(.*?)\n\s*Keywords:", dx_full, re.S)
    if m:
        n_words = len(m.group(1).split())
        good = 150 <= n_words <= 200
        print(f"{'✅' if good else '❌'} 초록 {n_words} 단어 (규정 150~200)")
        ok = ok and good

    bad_font = {f for p in ps for f in fonts(p)} - {"Times New Roman"}
    print(f"{'✅' if not bad_font else '❌'} 본문 글꼴 Times New Roman 단일"
          f"{'' if not bad_font else f' — 이물질: {bad_font}'}")
    ok = ok and not bad_font

    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
